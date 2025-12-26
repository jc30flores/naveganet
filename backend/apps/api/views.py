from rest_framework import viewsets
from rest_framework.decorators import api_view, parser_classes, action
from rest_framework.generics import ListAPIView
from rest_framework.parsers import JSONParser
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.views import APIView
from django.http import JsonResponse, HttpResponse, FileResponse
from django.db.models import (
    Q,
    F,
    Max,
    Min,
    Sum,
    DecimalField,
    ExpressionWrapper,
    Count,
    Exists,
    OuterRef,
    Subquery,
    Case,
    When,
    Value,
    Prefetch,
    Func,
)
from django.contrib.postgres.aggregates import StringAgg
from django.db.models.functions import Coalesce, TruncDay, TruncMonth, TruncYear, Lower
from django.db import models as dj_models
from datetime import timedelta, date, datetime, time as dt_time
from collections import defaultdict
import calendar
from django.db import DataError, IntegrityError, transaction, connection
from django.utils import timezone
from django.core.cache import cache
from django.core.paginator import Paginator
from django.conf import settings
import time
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
from io import BytesIO
import math
import unicodedata
from typing import List

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from . import models, serializers
from .db_state import has_unaccent, has_unaccent_wrapper


class Unaccent(Func):
    function = "UNACCENT"
    output_field = dj_models.TextField()


class ImmutableUnaccent(Func):
    function = "public.unaccent_immutable"
    output_field = dj_models.TextField()


def _normalized_unaccent(*field_names):
    expressions = [F(name) for name in field_names]
    expressions.append(Value("", output_field=dj_models.TextField()))
    base = Coalesce(*expressions, output_field=dj_models.TextField())
    if has_unaccent_wrapper():
        return Lower(ImmutableUnaccent(base))
    if has_unaccent():
        return Lower(Unaccent(base))
    return Lower(base)


def _strip_accents(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    return "".join(ch for ch in normalized if not unicodedata.combining(ch))
from .utils.security import constant_time_compare


PAYMENT_METHOD_MAP = {
    "CASH": "efectivo",
    "CARD": "tarjeta",
    "TRANSFER": "transferencia",
}

MONTH_NAMES_ES = [
    "",
    "Enero",
    "Febrero",
    "Marzo",
    "Abril",
    "Mayo",
    "Junio",
    "Julio",
    "Agosto",
    "Septiembre",
    "Octubre",
    "Noviembre",
    "Diciembre",
]


def _parse_date(value, today):
    if not value:
        return None
    d = date.fromisoformat(value)
    return today if d > today else d


def _infer_mode(start, end, today):
    if not start and not end:
        return "all"
    if start == today and end == today:
        return "daily"
    first = date(today.year, today.month, 1)
    if start == first and end == today:
        return "monthly"
    quinc_start = first if today.day <= 15 else date(today.year, today.month, 16)
    if start == quinc_start and end == today:
        return "quincenal"
    return "range"


def _range_label(mode, start, end, today):
    if mode == "daily":
        return f"Rango: Diario {today.strftime('%d/%m/%Y')}"
    if mode == "monthly":
        return f"Rango: {MONTH_NAMES_ES[today.month]}"
    if mode == "quincenal":
        month_name = MONTH_NAMES_ES[today.month]
        if today.day <= 15:
            return f"Rango: 1 - {today.day} / {month_name}"
        return f"Rango: 16 - {today.day} / {month_name}"
    if mode == "all":
        return "Rango: Todas Las Ventas"
    if start and end:
        return f"Rango: {start.isoformat()} - {end.isoformat()}"
    return "Rango:"


def _venta_product_names(venta_id):
    detalles = models.DetalleVenta.objects.filter(venta_id=venta_id).select_related("producto")
    nombres = []
    for d in detalles:
        nombres.append(d.producto_nombre_snapshot or (d.producto.nombre if d.producto else ""))
    return ", ".join(filter(None, nombres))


def health(request):
    return JsonResponse({'status': 'ok'})


@api_view(["GET"])
def ventas_historial(request):
    mode = request.query_params.get("mode", "daily")
    if mode not in {"daily", "quincenal", "monthly", "all", "range"}:
        return Response({"detail": "Invalid mode"}, status=400)

    now = timezone.localtime()
    today = now.date()

    start_date = end_date = None
    if mode == "range":
        start_str = request.query_params.get("start")
        end_str = request.query_params.get("end")
        start_date = date.fromisoformat(start_str) if start_str else None
        end_date = date.fromisoformat(end_str) if end_str else None
        if end_date and end_date > today:
            end_date = today
        if start_date and end_date and start_date > end_date:
            return Response({"detail": "Invalid range"}, status=400)
    elif mode == "daily":
        start_date = today
        end_date = today
    elif mode == "monthly":
        start_date = date(today.year, today.month, 1)
        end_date = today
    elif mode == "quincenal":
        start = 1 if today.day <= 15 else 16
        start_date = date(today.year, today.month, start)
        end_date = today
    # mode "all" leaves start_date/end_date as None (no filter)

    qs = models.Ventas.objects.select_related("cliente")

    if start_date:
        qs = qs.filter(fecha__date__gte=start_date)
    if end_date:
        qs = qs.filter(fecha__date__lte=end_date)

    q = request.query_params.get("q")
    if q:
        if q.isdigit():
            qs = qs.filter(Q(cliente__nombre__icontains=q) | Q(id=int(q)))
        else:
            qs = qs.filter(cliente__nombre__icontains=q)

    qs = qs.order_by("-fecha")

    page = int(request.query_params.get("page", 1))
    page_size = int(request.query_params.get("page_size", 30))
    if page_size > 30:
        page_size = 30

    paginator = Paginator(qs, page_size)
    page_obj = paginator.get_page(page)

    serializer = serializers.VentaListSerializer(page_obj.object_list, many=True)
    return Response(
        {
            "results": serializer.data,
            "page": page_obj.number,
            "page_size": page_size,
            "total_pages": paginator.num_pages,
            "total_count": paginator.count,
            "has_next": page_obj.has_next(),
            "has_prev": page_obj.has_previous(),
        }
    )


@api_view(["GET"])
def historial_ventas(request):
    try:
        mode = request.query_params.get("mode", "daily")
        if mode not in {"daily", "quincenal", "monthly", "all", "range"}:
            return Response({"detail": "Invalid mode"}, status=400)

        now = timezone.localtime()
        today = now.date()

        start_date = end_date = None
        if mode == "range":
            start_str = request.query_params.get("start")
            end_str = request.query_params.get("end")
            start_date = date.fromisoformat(start_str) if start_str else None
            end_date = date.fromisoformat(end_str) if end_str else None
            if end_date and end_date > today:
                end_date = today
            if start_date and end_date and start_date > end_date:
                return Response({"detail": "Invalid range"}, status=400)
        elif mode == "daily":
            start_date = today
            end_date = today
        elif mode == "monthly":
            start_date = date(today.year, today.month, 1)
            end_date = today
        elif mode == "quincenal":
            start = 1 if today.day <= 15 else 16
            start_date = date(today.year, today.month, start)
            end_date = today

        search = (request.query_params.get("q") or "").strip()

        venta_credit_exists = models.CreditosHistorialCompras.objects.filter(venta_id=OuterRef("pk"))

        contado_qs = (
            models.Ventas.objects.select_related("cliente")
            .annotate(tiene_credito=Exists(venta_credit_exists))
            .filter(tiene_credito=False)
        )
        if start_date:
            contado_qs = contado_qs.filter(fecha__date__gte=start_date)
        if end_date:
            contado_qs = contado_qs.filter(fecha__date__lte=end_date)
        if search:
            if search.isdigit():
                contado_qs = contado_qs.filter(
                    Q(id=int(search)) | Q(cliente__nombre__icontains=search)
                )
            else:
                contado_qs = contado_qs.filter(cliente__nombre__icontains=search)

        historial_subquery = models.CreditosHistorialCompras.objects.filter(
            credito_id=OuterRef("credito_id")
        ).order_by("fecha")

        abonos_qs = (
            models.PagosCredito.objects.select_related("credito__cliente")
            .annotate(venta_ref=Subquery(historial_subquery.values("venta_id")[:1]))
            .filter(venta_ref__isnull=False)
        )
        if start_date:
            abonos_qs = abonos_qs.filter(fecha__date__gte=start_date)
        if end_date:
            abonos_qs = abonos_qs.filter(fecha__date__lte=end_date)
        if search:
            if search.isdigit():
                abonos_qs = abonos_qs.filter(
                    Q(venta_ref=int(search))
                    | Q(credito__cliente__nombre__icontains=search)
                )
            else:
                abonos_qs = abonos_qs.filter(credito__cliente__nombre__icontains=search)

        try:
            page = int(request.query_params.get("page", 1))
        except (TypeError, ValueError):
            page = 1
        if page < 1:
            page = 1

        try:
            page_size = int(request.query_params.get("page_size", 30))
        except (TypeError, ValueError):
            page_size = 30
        page_size = max(1, min(page_size, 100))

        fetch_limit = page * page_size
        contado_items = list(contado_qs.order_by("-fecha")[:fetch_limit])
        abono_items = list(abonos_qs.order_by("-fecha")[:fetch_limit])

        devoluciones_qs = models.Devoluciones.objects.filter(venta__isnull=False)
        if start_date:
            devoluciones_qs = devoluciones_qs.filter(fecha__date__gte=start_date)
        if end_date:
            devoluciones_qs = devoluciones_qs.filter(fecha__date__lte=end_date)
        if search:
            if search.isdigit():
                devoluciones_qs = devoluciones_qs.filter(
                    Q(venta_id=int(search)) | Q(venta__cliente__nombre__icontains=search)
                )
            else:
                devoluciones_qs = devoluciones_qs.filter(
                    venta__cliente__nombre__icontains=search
                )

        devoluciones_total = (
            devoluciones_qs.values("venta_id", "fecha").distinct().count()
        )

        devoluciones_rows = list(
            devoluciones_qs
            .values(
                "venta_id",
                "fecha",
                "venta__cliente__id",
                "venta__cliente__nombre",
                "venta__cliente__razon_social",
                "venta__cliente__nombre_comercial",
                "venta__cliente__nit",
                "venta__cliente__telefono",
                "venta__documento_numero",
            )
            .annotate(
                total_refund=Coalesce(Sum("total"), Decimal("0")),
                ingreso_total=Coalesce(Sum("ingreso_afectado"), Decimal("0")),
                notas=StringAgg(
                    "motivo",
                    delimiter=" | ",
                    filter=~Q(motivo__isnull=True) & ~Q(motivo__exact=""),
                ),
                first_id=Min("id"),
            )
            .order_by("-fecha")[:fetch_limit]
        )

        contado_total = contado_qs.count()
        abono_total = abonos_qs.count()
        total_count = contado_total + abono_total + devoluciones_total

        combined = []
        for venta in contado_items:
            cliente = "Cliente General"
            if venta.cliente_id and venta.cliente:
                cliente = venta.cliente.nombre
            combined.append(
                {
                    "id": venta.id,
                    "fecha": venta.fecha,
                    "monto": float(venta.total),
                    "venta_id": venta.id,
                    "cliente": cliente,
                    "tipo": "VENTA_CONTADO",
                    "nota": None,
                    "ingreso_afectado": float(venta.total),
                    "venta_numero": venta.documento_numero,
                }
            )

        for abono in abono_items:
            credito = getattr(abono, "credito", None)
            cliente_obj = getattr(credito, "cliente", None)
            cliente = getattr(cliente_obj, "nombre", "Cliente General")
            combined.append(
                {
                    "id": abono.id,
                    "fecha": abono.fecha,
                    "monto": float(abono.monto),
                    "venta_id": abono.venta_ref,
                    "cliente": cliente,
                    "tipo": "ABONO",
                    "nota": None,
                    "ingreso_afectado": float(abono.monto),
                    "venta_numero": None,
                }
            )

        for row in devoluciones_rows:
            cliente_nombre = (
                row.get("venta__cliente__nombre")
                or row.get("venta__cliente__razon_social")
                or row.get("venta__cliente__nombre_comercial")
                or "Cliente General"
            )
            notas = row.get("notas") or None
            total_refund = row.get("total_refund") or Decimal("0")
            ingreso_total = row.get("ingreso_total") or Decimal("0")
            combined.append(
                {
                    "id": row.get("first_id"),
                    "fecha": row["fecha"],
                    "monto": -float(total_refund),
                    "venta_id": row["venta_id"],
                    "cliente": cliente_nombre,
                    "tipo": "DEVOLUCION",
                    "nota": notas,
                    "ingreso_afectado": float(ingreso_total),
                    "venta_numero": row.get("venta__documento_numero"),
                }
            )

        combined.sort(key=lambda x: x["fecha"], reverse=True)
        start = (page - 1) * page_size
        end = start + page_size
        page_rows = combined[start:end]

        results = [
            {
                **row,
                "fecha": timezone.localtime(row["fecha"]).isoformat(),
            }
            for row in page_rows
        ]

        total_pages = (total_count + page_size - 1) // page_size if total_count else 1
        return Response(
            {
                "count": total_count,
                "page": page,
                "page_size": page_size,
                "total_pages": total_pages,
                "results": results,
            }
        )
    except Exception as e:
        import traceback
        error_detail = str(e)
        traceback.print_exc()
        return Response({"detail": f"Server error: {error_detail}"}, status=500)


@api_view(["GET"])
def ventas_search(request):
    raw_term = (request.query_params.get("q") or "").strip()
    try:
        limit = int(request.query_params.get("limit", 10))
    except (TypeError, ValueError):
        limit = 10
    limit = max(1, min(limit, 50))

    try:
        page = int(request.query_params.get("page", 1))
    except (TypeError, ValueError):
        page = 1
    if page < 1:
        page = 1

    branch = (request.query_params.get("branch") or "").strip()
    cleaned = raw_term.lstrip("#").strip()
    normalized = " ".join(cleaned.split()).lower()
    supports_unaccent = has_unaccent()
    normalized_for_lookup = (
        _strip_accents(normalized) if supports_unaccent else normalized
    )

    if not normalized:
        return Response(
            {
                "results": [],
                "count": 0,
                "page": page,
                "page_size": limit,
                "total_pages": 0,
                "has_next": False,
                "has_prev": False,
            }
        )

    ventas_qs = (
        models.Ventas.objects.select_related("cliente")
        .annotate(
            es_credito=Exists(
                models.CreditosHistorialCompras.objects.filter(venta_id=OuterRef("pk"))
            )
        )
        .exclude(estado__iexact="anulada")
        .exclude(estado__iexact="cancelada")
    )

    if branch:
        field_names = {field.attname for field in models.Ventas._meta.concrete_fields}
        if "branch_id" in field_names:
            ventas_qs = ventas_qs.filter(branch_id=branch)
        elif "sucursal_id" in field_names:
            ventas_qs = ventas_qs.filter(sucursal_id=branch)
        elif "documento_serie" in field_names:
            ventas_qs = ventas_qs.filter(documento_serie__iexact=branch)

    ventas_qs = ventas_qs.annotate(
        numero_normalized=_normalized_unaccent("documento_numero"),
        cliente_nombre_normalized=_normalized_unaccent("cliente__nombre"),
        cliente_razon_normalized=_normalized_unaccent("cliente__razon_social"),
        cliente_comercial_normalized=_normalized_unaccent("cliente__nombre_comercial"),
        cliente_nit_normalized=_normalized_unaccent("cliente__nit"),
        cliente_tel_normalized=_normalized_unaccent("cliente__telefono"),
        cliente_dui_normalized=_normalized_unaccent("cliente__dui"),
    )

    detalle_match = models.DetalleVenta.objects.filter(venta_id=OuterRef("pk")).annotate(
        prod_codigo_snapshot_norm=_normalized_unaccent("producto_codigo_snapshot"),
        prod_codigo_live_norm=_normalized_unaccent("producto__codigo"),
        prod_nombre_snapshot_norm=_normalized_unaccent("producto_nombre_snapshot"),
        prod_nombre_live_norm=_normalized_unaccent("producto__nombre"),
    )
    detalle_match = detalle_match.filter(
        Q(prod_codigo_snapshot_norm__contains=normalized_for_lookup)
        | Q(prod_codigo_live_norm__contains=normalized_for_lookup)
        | Q(prod_nombre_snapshot_norm__contains=normalized_for_lookup)
        | Q(prod_nombre_live_norm__contains=normalized_for_lookup)
    )

    ventas_qs = ventas_qs.annotate(producto_match=Exists(detalle_match))

    filters = Q(numero_normalized__contains=normalized_for_lookup)
    filters |= Q(cliente_nombre_normalized__contains=normalized_for_lookup)
    filters |= Q(cliente_razon_normalized__contains=normalized_for_lookup)
    filters |= Q(cliente_comercial_normalized__contains=normalized_for_lookup)
    filters |= Q(cliente_nit_normalized__contains=normalized_for_lookup)
    filters |= Q(cliente_tel_normalized__contains=normalized_for_lookup)
    filters |= Q(cliente_dui_normalized__contains=normalized_for_lookup)

    if normalized.isdigit():
        try:
            filters |= Q(id=int(normalized))
        except (TypeError, ValueError):
            pass
        filters |= Q(numero_normalized__contains=normalized_for_lookup)

    ventas_qs = ventas_qs.filter(filters | Q(producto_match=True)).distinct()

    ventas_qs = ventas_qs.order_by("-fecha", "-id")
    total_count = ventas_qs.count()
    if total_count == 0:
        return Response(
            {
                "results": [],
                "count": 0,
                "page": 1,
                "page_size": limit,
                "total_pages": 0,
                "has_next": False,
                "has_prev": False,
            }
        )

    total_pages = (total_count + limit - 1) // limit
    page_number = min(page, total_pages) if total_pages else 1
    if page_number < 1:
        page_number = 1
    offset = (page_number - 1) * limit

    detalle_qs = models.DetalleVenta.objects.select_related("producto")
    ventas = list(
        ventas_qs.prefetch_related(
            Prefetch("detalles", queryset=detalle_qs, to_attr="prefetched_detalles")
        )[offset : offset + limit]
    )

    def _decimal_to_number(value):
        if value is None:
            return 0
        if isinstance(value, Decimal):
            return float(value)
        return float(Decimal(value))

    results = []
    for venta in ventas:
        detalles = getattr(venta, "prefetched_detalles", [])
        subtotal = Decimal("0")
        items = []
        total_returnable = Decimal("0")
        for det in detalles:
            subtotal += det.subtotal or Decimal("0")
            devuelto = det.devuelto or Decimal("0")
            restante = det.cantidad - devuelto
            if restante < Decimal("0"):
                restante = Decimal("0")
            total_returnable += restante
            codigo = det.producto_codigo_snapshot or getattr(det.producto, "codigo", "")
            if not codigo and getattr(det.producto, "codigo", None):
                codigo = det.producto.codigo
            nombre = det.producto_nombre_snapshot or getattr(det.producto, "nombre", "")
            if not nombre and getattr(det.producto, "nombre", None):
                nombre = det.producto.nombre
            condicion = (
                det.producto_condicion_snapshot or getattr(det.producto, "condicion", "") or ""
            )
            condicion_lower = condicion.lower() if condicion else ""
            if condicion_lower in {"nuevo", "new"}:
                condicion_label = "nuevo"
            elif condicion_lower in {"usado", "used"}:
                condicion_label = "usado"
            else:
                condicion_label = condicion_lower or None
            items.append(
                {
                    "detalle_id": det.id,
                    "producto_id": det.producto_id,
                    "codigo": codigo,
                    "nombre": nombre,
                    "cantidad": _decimal_to_number(det.cantidad),
                    "devuelto": _decimal_to_number(devuelto),
                    "disponible": _decimal_to_number(restante),
                    "precio": _decimal_to_number(det.precio_unitario),
                    "subtotal": _decimal_to_number(det.subtotal),
                    "nuevo_usado": condicion_label,
                    "condicion": condicion_lower or None,
                }
            )

        cliente = venta.cliente
        cliente_data = None
        if cliente:
            nombre_cliente = (
                cliente.nombre
                or cliente.razon_social
                or cliente.nombre_comercial
                or "Cliente general"
            )
            cliente_data = {
                "id": cliente.id,
                "nombre": nombre_cliente,
                "nit": cliente.nit,
                "telefono": cliente.telefono,
            }

        results.append(
            {
                "id": venta.id,
                "numero": venta.documento_numero or str(venta.id),
                "fecha": venta.fecha,
                "cliente": cliente_data,
                "tipo": "credito" if getattr(venta, "es_credito", False) else "contado",
                "subtotal": _decimal_to_number(subtotal),
                "impuestos": _decimal_to_number(venta.iva_monto),
                "total": _decimal_to_number(venta.total),
                "items": items,
                "returnable_total": _decimal_to_number(total_returnable),
            }
        )

    has_next = page_number < total_pages
    has_prev = page_number > 1

    return Response(
        {
            "results": results,
            "count": total_count,
            "page": page_number,
            "page_size": limit,
            "total_pages": total_pages,
            "has_next": has_next,
            "has_prev": has_prev,
        }
    )


@api_view(["GET"])
def ventas_export(request):
    fmt = request.query_params.get("format")
    if fmt not in {"pdf", "xlsx", "docx"}:
        return Response({"detail": "Invalid format"}, status=400)

    mode = request.query_params.get("mode")
    if mode and mode not in {"daily", "monthly", "quincenal", "all", "range"}:
        return Response({"detail": "Invalid mode"}, status=400)

    now = timezone.localtime()
    today = now.date()

    start_date = _parse_date(request.query_params.get("start"), today)
    end_date = _parse_date(request.query_params.get("end"), today)
    if start_date and end_date and start_date > end_date:
        return Response({"detail": "Invalid range"}, status=400)

    mode = mode or _infer_mode(start_date, end_date, today)
    label = _range_label(mode, start_date, end_date, today)

    qs = models.Ventas.objects.select_related("cliente").order_by("fecha", "id")
    if start_date:
        qs = qs.filter(fecha__date__gte=start_date)
    if end_date:
        qs = qs.filter(fecha__date__lte=end_date)

    entries = []
    grand = Decimal("0")
    for v in qs:
        cliente = v.cliente.nombre if v.cliente else "Cliente General"
        local_dt = timezone.localtime(v.fecha)
        fecha_str = local_dt.strftime("%d/%m/%Y %I:%M %p")
        productos = _venta_product_names(v.id)
        total = v.total or Decimal("0")
        grand += total
        entries.append(
            {
                "venta_id": v.id,
                "cliente": cliente,
                "fecha_dt": local_dt,
                "fecha_str": fecha_str,
                "descripcion": productos,
                "total": total,
                "tipo": "VENTA",
            }
        )

    devoluciones_qs = models.Devoluciones.objects.filter(venta__isnull=False)
    if start_date:
        devoluciones_qs = devoluciones_qs.filter(fecha__date__gte=start_date)
    if end_date:
        devoluciones_qs = devoluciones_qs.filter(fecha__date__lte=end_date)

    devoluciones_rows = (
        devoluciones_qs.values(
            "venta_id",
            "fecha",
            "venta__cliente__nombre",
            "venta__cliente__razon_social",
            "venta__cliente__nombre_comercial",
            "venta__documento_numero",
        )
        .annotate(
            total_refund=Coalesce(Sum("total"), Decimal("0")),
            notas=StringAgg(
                "motivo",
                delimiter=" | ",
                filter=~Q(motivo__isnull=True) & ~Q(motivo__exact=""),
            ),
        )
        .order_by("fecha", "venta_id")
    )

    for row in devoluciones_rows:
        local_dt = timezone.localtime(row["fecha"])
        fecha_str = local_dt.strftime("%d/%m/%Y %I:%M %p")
        cliente = (
            row.get("venta__cliente__nombre")
            or row.get("venta__cliente__razon_social")
            or row.get("venta__cliente__nombre_comercial")
            or "Cliente General"
        )
        numero = row.get("venta__documento_numero") or str(row["venta_id"])
        motivo = row.get("notas")
        descripcion = f"Devolución venta #{numero}"
        if motivo:
            descripcion = f"{descripcion} · {motivo}"
        total_refund = row.get("total_refund") or Decimal("0")
        grand -= total_refund
        entries.append(
            {
                "venta_id": row["venta_id"],
                "cliente": cliente,
                "fecha_dt": local_dt,
                "fecha_str": fecha_str,
                "descripcion": descripcion,
                "total": -total_refund,
                "tipo": "DEVOLUCION",
            }
        )

    entries.sort(
        key=lambda item: (
            item["fecha_dt"],
            item.get("venta_id") or 0,
            0 if item["tipo"] == "VENTA" else 1,
        )
    )
    count = len(entries)
    dataset = [
        [
            entry["cliente"],
            entry["fecha_str"],
            entry["descripcion"],
            entry["total"],
        ]
        for entry in entries
    ]
    dataset.append(["TOTAL", "", "", grand])

    if fmt == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
        except Exception:
            return Response({"detail": "Word export not available"}, status=500)
        doc = Document()
        title = doc.add_paragraph()
        title.alignment = 1
        title_run = title.add_run("Historial de Ventas")
        title_run.font.size = Pt(22)
        rng_p = doc.add_paragraph()
        rng_p.alignment = 1
        rng_run = rng_p.add_run(label)
        rng_run.font.size = Pt(12)
        table = doc.add_table(rows=1, cols=4)
        headers = ["Cliente", "Fecha", "Productos", "Total"]
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for r in hdr_cells[i].paragraphs[0].runs:
                r.bold = True
        for row in dataset:
            cells = table.add_row().cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            cells[2].text = row[2]
            cells[3].text = f"${row[3]:,.2f}"
            if row[0] == "TOTAL":
                for cell in cells:
                    for r in cell.paragraphs[0].runs:
                        r.bold = True
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    elif fmt == "xlsx":
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment
            from openpyxl.utils import get_column_letter
        except Exception:
            return Response({"detail": "Excel export not available"}, status=500)
        wb = Workbook()
        ws = wb.active
        ws.title = "Historial"
        ws.merge_cells("A1:D1")
        ws["A1"].value = "Historial de Ventas"
        ws["A1"].font = Font(size=22, bold=True)
        ws["A1"].alignment = Alignment(horizontal="center")
        ws.merge_cells("A2:D2")
        ws["A2"].value = label
        ws["A2"].font = Font(size=12)
        ws["A2"].alignment = Alignment(horizontal="center")
        headers = ["Cliente", "Fecha", "Productos", "Total"]
        ws.append(headers)
        for cell in ws[3]:
            cell.font = Font(bold=True)
        for row in dataset:
            ws.append([row[0], row[1], row[2], float(row[3])])
        last = len(dataset) + 3
        for cell in ws[last]:
            cell.font = Font(bold=True)
        widths = [28, 20, 48, 14]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w
        for r in range(4, last + 1):
            ws.cell(row=r, column=4).number_format = u'"$"#,##0.00'
            ws.cell(row=r, column=2).alignment = Alignment(horizontal="left")
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        resp = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    else:  # pdf
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except Exception:
            return Response({"detail": "PDF export not available"}, status=500)

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=letter,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            name="Title",
            parent=styles["Title"],
            alignment=1,
            fontSize=22,
        )
        range_style = ParagraphStyle(
            name="Range",
            parent=styles["Normal"],
            alignment=1,
            fontSize=12,
        )
        body_style = ParagraphStyle(
            name="Body",
            parent=styles["BodyText"],
            fontSize=10,
            leading=12,
        )
        body_style.wordWrap = "CJK"
        right_style = ParagraphStyle(
            name="Right",
            parent=body_style,
            alignment=2,
        )
        right_style.wordWrap = "CJK"

        headers = ["Cliente", "Fecha", "Productos", "Total"]
        data = [headers]
        for entry in entries:
            data.append(
                [
                    Paragraph(entry.get("cliente") or "", body_style),
                    Paragraph(entry.get("fecha_str") or "", body_style),
                    Paragraph(entry.get("descripcion") or "", body_style),
                    Paragraph(f"${float(entry.get('total', 0)):,.2f}", right_style),
                ]
            )
        data.append(
            [
                Paragraph("TOTAL", body_style),
                Paragraph("", body_style),
                Paragraph("", body_style),
                Paragraph(f"${float(grand):,.2f}", right_style),
            ]
        )

        col_widths = [140, 110, 220, 60]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("FONT", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
                    ("ALIGN", (3, 1), (3, -1), "RIGHT"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.gray),
                ]
            )
        )

        elements = [
            Paragraph("Historial de Ventas", title_style),
            Spacer(1, 12),
            Paragraph(label.replace("Rango:", "Rango:"), range_style),
            Spacer(1, 12),
            table,
        ]
        doc.build(elements)
        resp = HttpResponse(buf.getvalue(), content_type="application/pdf")

    fname = timezone.localtime().strftime("historial_de_venta_%Y%m%d_%H%M%S") + f".{fmt}"
    resp["Content-Disposition"] = f'attachment; filename="{fname}"'
    resp["X-Filename"] = fname
    resp["X-Row-Count"] = str(count)
    return resp


class ReporteExportInventarioView(APIView):
    permission_classes = [IsAuthenticated]

    COLUMN_KEYS: List[str] = [
        "codigo",
        "nombre",
        "categoria",
        "condicion",
        "precio",
        "costo",
        "stock",
        "stock_minimo",
    ]
    REPORT_HEADERS: List[str] = [
        "Código",
        "Producto",
        "Categoría",
        "Condición",
        "Precio",
        "Costo",
        "Stock",
        "Stock mínimo",
    ]
    RENAME_MAP = dict(zip(COLUMN_KEYS, REPORT_HEADERS))
    currency_fields = {"precio", "costo"}
    integer_fields = {"stock", "stock_minimo"}
    format_content_types = {
        "pdf": "application/pdf",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }
    condition_aliases = {
        "todos": "todos",
        "all": "todos",
        "nuevos": "nuevos",
        "new": "nuevos",
        "usados": "usados",
        "used": "usados",
    }
    condition_filters = {
        "todos": None,
        "nuevos": "NEW",
        "usados": "USED",
    }
    condition_labels = {
        "todos": "Todos",
        "nuevos": "Nuevos",
        "usados": "Usados",
    }
    slug_map = {
        "todos": "todos",
        "nuevos": "nuevos",
        "usados": "usados",
    }

    @staticmethod
    def _format_currency(value):
        if value is None:
            return Decimal("0.00")
        if isinstance(value, Decimal):
            return value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        try:
            return Decimal(str(value)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        except (InvalidOperation, TypeError, ValueError):
            return Decimal("0.00")

    @staticmethod
    def _format_int(value):
        try:
            return int(value)
        except (TypeError, ValueError):
            return 0

    def _build_dataset(self, queryset):
        rows: List[dict] = []
        for producto in queryset:
            categoria_nombre = producto.categoria.nombre if producto.categoria_id else ""
            condicion_raw = (producto.condicion or "").lower()
            if condicion_raw == "new":
                condicion_display = "Nuevo"
            elif condicion_raw == "used":
                condicion_display = "Usado"
            else:
                condicion_display = producto.condicion or ""

            rows.append(
                {
                    "codigo": producto.codigo or "",
                    "nombre": producto.nombre or "",
                    "categoria": categoria_nombre or "",
                    "condicion": condicion_display,
                    "precio": self._format_currency(producto.precio),
                    "costo": self._format_currency(producto.costo),
                    "stock": self._format_int(producto.stock),
                    "stock_minimo": self._format_int(producto.stock_minimo),
                }
            )
        return rows

    def _normalize_dataset(self, dataset):
        if dataset is None:
            raw_rows = []
        elif isinstance(dataset, pd.DataFrame):
            df = dataset.copy()
            for key in self.COLUMN_KEYS:
                if key not in df.columns:
                    df[key] = None
            df = df[self.COLUMN_KEYS]
            raw_rows = df.to_dict(orient="records")
        elif isinstance(dataset, dict):
            raw_rows = [dataset]
        else:
            raw_rows = list(dataset)

        normalized_rows: List[dict] = []
        alias_map = {"producto": "nombre"}
        for raw in raw_rows:
            if isinstance(raw, pd.Series):
                raw = raw.to_dict()
            elif hasattr(raw, "_asdict"):
                raw = raw._asdict()
            elif not isinstance(raw, dict):
                raw = {key: getattr(raw, key, None) for key in self.COLUMN_KEYS}

            row: dict = {}
            for key in self.COLUMN_KEYS:
                value = None
                if isinstance(raw, dict):
                    if key in raw:
                        value = raw.get(key)
                    elif self.RENAME_MAP[key] in raw:
                        value = raw.get(self.RENAME_MAP[key])
                    else:
                        alias_key = alias_map.get(key)
                        if alias_key and alias_key in raw:
                            value = raw.get(alias_key)
                        else:
                            lower_key = key.lower()
                            value = raw.get(lower_key)
                if key in self.currency_fields:
                    row[key] = self._format_currency(value)
                elif key in self.integer_fields:
                    row[key] = self._format_int(value)
                else:
                    if value in (None, "None"):
                        row[key] = ""
                    else:
                        row[key] = str(value) if not isinstance(value, str) else value
            normalized_rows.append(row)

        return normalized_rows

    def _render_pdf(self, rows, condition_label):
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=landscape(A4),
            leftMargin=24,
            rightMargin=24,
            topMargin=24,
            bottomMargin=24,
        )
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            name="InventoryTitle",
            parent=styles["Title"],
            alignment=1,
            fontSize=22,
        )
        subtitle_style = ParagraphStyle(
            name="InventorySubtitle",
            parent=styles["Normal"],
            alignment=1,
            fontSize=12,
        )
        body_style = ParagraphStyle(
            name="InventoryBody",
            parent=styles["BodyText"],
            fontSize=9,
            leading=11,
        )
        body_style.wordWrap = "CJK"
        right_style = ParagraphStyle(
            name="InventoryRight",
            parent=body_style,
            alignment=2,
        )
        right_style.wordWrap = "CJK"

        def format_cell(key, value):
            text = ""
            if key in self.currency_fields:
                if str(value) not in {"", "None"}:
                    try:
                        text = f"${float(value):,.2f}"
                    except (TypeError, ValueError):
                        text = str(value or "")
            else:
                text = str(value or "")
            style = right_style if key in self.currency_fields else body_style
            return Paragraph(text, style)

        data_rows = [
            [format_cell(key, row.get(key, "")) for key in self.COLUMN_KEYS]
            for row in rows
        ]
        if not data_rows:
            data_rows.append(
                [Paragraph("", body_style) for _ in self.COLUMN_KEYS]
            )

        table_data = [self.REPORT_HEADERS] + data_rows

        col_widths = [70, 220, 110, 85, 80, 80, 60, 90]
        table = Table(table_data, repeatRows=1, colWidths=col_widths)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e6e6e6")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                    ("ALIGN", (4, 1), (5, -1), "RIGHT"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.gray),
                    ("WORDWRAP", (0, 0), (-1, -1), True),
                ]
            )
        )

        elements = [
            Paragraph("Reporte de Inventario", title_style),
            Spacer(1, 12),
            Paragraph(f"Condición: {condition_label}", subtitle_style),
            Spacer(1, 18),
            table,
        ]
        doc.build(elements)
        buffer.seek(0)
        return buffer

    def _render_xlsx(self, rows):
        df = (
            pd.DataFrame(rows, columns=self.COLUMN_KEYS)
            if rows
            else pd.DataFrame(columns=self.COLUMN_KEYS)
        )
        df = df.rename(columns=self.RENAME_MAP)

        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
            df.to_excel(writer, sheet_name="Inventario", index=False)
            workbook = writer.book
            worksheet = writer.sheets["Inventario"]

            from openpyxl.styles import Alignment, Font
            from openpyxl.utils import get_column_letter

            for cell in worksheet[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")

            if worksheet.max_row >= 1 and worksheet.max_column >= 1:
                worksheet.auto_filter.ref = worksheet.dimensions

            currency_columns = {"Precio", "Costo"}
            numeric_columns = {"Stock", "Stock mínimo"}

            for col_idx in range(1, worksheet.max_column + 1):
                header = worksheet.cell(row=1, column=col_idx).value
                max_length = len(str(header)) if header else 0
                for row_idx in range(2, worksheet.max_row + 1):
                    cell = worksheet.cell(row=row_idx, column=col_idx)
                    if cell.value is None:
                        cell.value = "" if header not in (currency_columns | numeric_columns) else 0
                    display_value = cell.value
                    if header in currency_columns:
                        cell.number_format = '"$"#,##0.00'
                        cell.alignment = Alignment(horizontal="right")
                        try:
                            cell.value = float(display_value)
                        except (TypeError, ValueError):
                            cell.value = 0.0
                    elif header in numeric_columns:
                        cell.alignment = Alignment(horizontal="right")
                    length = len(str(cell.value))
                    if length > max_length:
                        max_length = length
                adjusted_width = min(max(max_length + 2, 12), 40)
                worksheet.column_dimensions[get_column_letter(col_idx)].width = adjusted_width

        buffer.seek(0)
        return buffer

    def _render_docx(self, rows, header_labels, column_keys, condition_label):
        document = Document()
        section = document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        title = document.add_paragraph()
        title.alignment = 1
        title_run = title.add_run("Reporte de Inventario")
        title_run.bold = True
        title_run.font.size = Pt(22)

        subtitle = document.add_paragraph()
        subtitle.alignment = 1
        subtitle_run = subtitle.add_run(f"Condición: {condition_label}")
        subtitle_run.font.size = Pt(12)

        table = document.add_table(rows=1, cols=len(header_labels))
        table.style = "Table Grid"
        table.autofit = False
        header_row = table.rows[0]
        for idx, header in enumerate(header_labels):
            cell = header_row.cells[idx]
            cell.text = header
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            if cell.paragraphs:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        data_rows = rows if rows else []
        if not data_rows:
            empty_cells = table.add_row().cells
            for idx in range(len(column_keys)):
                empty_cells[idx].text = ""
        for row in data_rows:
            row_cells = table.add_row().cells
            for idx, key in enumerate(column_keys):
                value = row.get(key, "") if row else ""
                if key in self.currency_fields:
                    if str(value) not in {"", "None"}:
                        try:
                            text = f"${float(value):,.2f}"
                        except (TypeError, ValueError):
                            text = str(value)
                    else:
                        text = ""
                else:
                    text = "" if value in (None, "None") else str(value)
                cell = row_cells[idx]
                cell.text = text
                if cell.paragraphs:
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT
                        if key in self.currency_fields
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )

        width_points = [70, 220, 110, 85, 80, 80, 60, 90]
        widths = [Inches(w / 72) for w in width_points]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def get(self, request):
        format_param = (request.query_params.get("format") or "pdf").lower()
        condicion_param = (request.query_params.get("condicion") or "todos").lower()

        if format_param == "docx":
            return Response(
                {"detail": "Formato DOCX no disponible para inventario."},
                status=400,
            )

        if format_param not in self.format_content_types:
            return Response({"detail": "Formato inválido"}, status=400)

        canonical_condition = self.condition_aliases.get(condicion_param)
        if canonical_condition is None:
            return Response({"detail": "Condición inválida"}, status=400)

        condition_filter = self.condition_filters[canonical_condition]
        productos_qs = (
            models.Productos.objects.filter(status="active")
            .select_related("categoria")
            .order_by("nombre")
        )
        if condition_filter:
            productos_qs = productos_qs.filter(condicion__iexact=condition_filter)

        dataset = self._build_dataset(productos_qs)
        rows = self._normalize_dataset(dataset)
        rows = [
            {key: (row or {}).get(key, "") for key in self.COLUMN_KEYS}
            for row in rows
        ]
        row_count = len(rows)
        condition_label = self.condition_labels[canonical_condition]
        slug = self.slug_map[canonical_condition]
        timestamp = timezone.localtime().strftime("%Y%m%d_%H%M%S")
        filename = f"inventario_{slug}_{timestamp}.{format_param}"

        if format_param == "pdf":
            buffer = self._render_pdf(rows, condition_label)
        elif format_param == "xlsx":
            buffer = self._render_xlsx(rows)

        response = FileResponse(
            buffer,
            as_attachment=True,
            filename=filename,
            content_type=self.format_content_types[format_param],
        )
        response["X-Row-Count"] = str(row_count)
        return response


@api_view(["GET"])
def ventas_total(request):
    now = timezone.localtime()
    today = now.date()

    start_str = request.query_params.get("start")
    end_str = request.query_params.get("end")

    start_date = date.fromisoformat(start_str) if start_str else None
    end_date = date.fromisoformat(end_str) if end_str else None
    if end_date and end_date > today:
        end_date = today
    if start_date and end_date and start_date > end_date:
        return Response({"detail": "Invalid range"}, status=400)

    cash_qs = models.Ventas.objects.filter(creditoshistorialcompras__isnull=True)
    if start_date:
        cash_qs = cash_qs.filter(fecha__date__gte=start_date)
    if end_date:
        cash_qs = cash_qs.filter(fecha__date__lte=end_date)
    cash_total = cash_qs.aggregate(total=Coalesce(Sum("total"), Decimal("0")))[
        "total"
    ] or Decimal("0")

    cash_refunds_qs = models.Devoluciones.objects.filter(
        venta__creditoshistorialcompras__isnull=True
    )
    if start_date:
        cash_refunds_qs = cash_refunds_qs.filter(fecha__date__gte=start_date)
    if end_date:
        cash_refunds_qs = cash_refunds_qs.filter(fecha__date__lte=end_date)
    cash_refunds = cash_refunds_qs.aggregate(
        total=Coalesce(Sum("ingreso_afectado"), Decimal("0"))
    )["total"] or Decimal("0")

    credit_qs = models.PagosCredito.objects.all()
    if start_date:
        credit_qs = credit_qs.filter(fecha__date__gte=start_date)
    if end_date:
        credit_qs = credit_qs.filter(fecha__date__lte=end_date)
    credit_total = credit_qs.aggregate(total=Coalesce(Sum("monto"), Decimal("0")))[
        "total"
    ] or Decimal("0")

    credit_refunds_qs = models.Devoluciones.objects.filter(
        venta__creditoshistorialcompras__isnull=False
    )
    if start_date:
        credit_refunds_qs = credit_refunds_qs.filter(fecha__date__gte=start_date)
    if end_date:
        credit_refunds_qs = credit_refunds_qs.filter(fecha__date__lte=end_date)
    credit_refunds = credit_refunds_qs.aggregate(
        total=Coalesce(Sum("ingreso_afectado"), Decimal("0"))
    )["total"] or Decimal("0")

    total = (cash_total - cash_refunds) + (credit_total - credit_refunds)

    return Response(
        {
            "total": total,
            "start": start_date.isoformat() if start_date else None,
            "end": end_date.isoformat() if end_date else None,
        }
    )


@api_view(["GET"])
def reportes_dashboard(request):
    section = request.GET.get("section", "all")
    if section == "new":
        prod_cond = "new"
        product_filter = {"condicion": "new"}
        detalle_filter = {"producto__condicion": "new"}
    elif section == "used":
        prod_cond = "used"
        product_filter = {"condicion": "used"}
        detalle_filter = {"producto__condicion": "used"}
    else:
        prod_cond = None
        product_filter = {}
        detalle_filter = {}

    now = timezone.localtime()
    today = now.date()
    tz = timezone.get_current_timezone()

    def _aware_datetime(d: date, t: dt_time = dt_time.min) -> datetime:
        naive = datetime.combine(d, t)
        if timezone.is_naive(naive):
            return timezone.make_aware(naive, tz)
        return naive.astimezone(tz)

    # Stats
    total_productos = models.Productos.objects.filter(**product_filter).count()
    start_today = _aware_datetime(today)
    end_today = start_today + timedelta(days=1)
    ventas_hoy = Decimal("0")

    creditos_qs = models.CreditosHistorialCompras.objects.all()
    if prod_cond:
        creditos_qs = creditos_qs.filter(
            venta__detalles__producto__condicion=prod_cond
        ).distinct()
    creditos_pendientes = creditos_qs.aggregate(
        total=Coalesce(Sum("saldo"), Decimal("0"))
    )["total"]

    valor_inventario = models.Productos.objects.filter(**product_filter).aggregate(
        total=Coalesce(
            Sum(
                ExpressionWrapper(
                    F("costo") * F("stock"),
                    output_field=DecimalField(max_digits=14, decimal_places=2),
                )
            ),
            Decimal("0"),
        )
    )["total"]

    devoluciones_qs = models.Devoluciones.objects.filter(
        fecha__year=today.year, fecha__month=today.month
    )
    if prod_cond:
        devoluciones_qs = devoluciones_qs.filter(producto__condicion=prod_cond)
    devoluciones_mensuales = devoluciones_qs.aggregate(
        total=Coalesce(Sum("total"), Decimal("0"))
    )["total"]

    stats = {
        "total_productos": total_productos,
        "ventas_hoy": ventas_hoy,
        "creditos_pendientes": creditos_pendientes,
        "valor_inventario": valor_inventario,
        "devoluciones_mensuales": devoluciones_mensuales,
    }

    # Sales chart data
    cost_expr = ExpressionWrapper(
        F("cantidad")
        * Coalesce("producto_costo_snapshot", F("producto__costo")),
        output_field=DecimalField(max_digits=14, decimal_places=2),
    )
    detalles_sales = models.DetalleVenta.objects.filter(**detalle_filter)
    detalles_profit = detalles_sales
    if section == "all":
        detalles_profit = detalles_sales.filter(producto__condicion="new")

    condition_new = Q(producto_condicion_snapshot__iexact="new") | Q(
        producto__condicion__iexact="new"
    )
    condition_used = Q(producto_condicion_snapshot__iexact="used") | Q(
        producto__condicion__iexact="used"
    )

    cash_details_base = models.DetalleVenta.objects.filter(
        venta__creditoshistorialcompras__isnull=True
    )

    credit_sale_lookup = models.CreditosHistorialCompras.objects.filter(
        venta_id=OuterRef("venta_id")
    ).order_by("fecha")

    credit_detail_rows = (
        models.DetalleVenta.objects.annotate(
            credito_ref=Subquery(credit_sale_lookup.values("credito_id")[:1])
        )
        .filter(credito_ref__isnull=False)
        .values("credito_ref")
        .annotate(
            new_total=Coalesce(
                Sum("subtotal", filter=condition_new),
                Decimal("0"),
            ),
            used_total=Coalesce(
                Sum("subtotal", filter=condition_used),
                Decimal("0"),
            ),
        )
    )

    credit_ratios = {}
    default_ratio = {"new": Decimal("1"), "used": Decimal("0")}
    for row in credit_detail_rows:
        new_total = row["new_total"] or Decimal("0")
        used_total = row["used_total"] or Decimal("0")
        base_total = new_total + used_total
        if base_total > 0:
            credit_ratios[row["credito_ref"]] = {
                "new": new_total / base_total,
                "used": used_total / base_total,
            }
        elif new_total > 0:
            credit_ratios[row["credito_ref"]] = {"new": Decimal("1"), "used": Decimal("0")}
        elif used_total > 0:
            credit_ratios[row["credito_ref"]] = {"new": Decimal("0"), "used": Decimal("1")}
        else:
            credit_ratios[row["credito_ref"]] = default_ratio

    def _refund_condition_filter(qs, condition):
        if condition == "new":
            return qs.filter(
                Q(producto_condicion_snapshot__iexact="new")
                | Q(producto__condicion__iexact="new")
            )
        if condition == "used":
            return qs.filter(
                Q(producto_condicion_snapshot__iexact="used")
                | Q(producto__condicion__iexact="used")
            )
        return qs

    def _aggregate_cash_period(qs, period_expr, start=None, end=None, condition=None):
        filtered = qs
        if start is not None:
            filtered = filtered.filter(venta__fecha__gte=start)
        if end is not None:
            filtered = filtered.filter(venta__fecha__lt=end)
        rows = (
            filtered.annotate(periodo=period_expr)
            .values("periodo")
            .annotate(total=Coalesce(Sum("subtotal"), Decimal("0")))
        )
        data = defaultdict(lambda: Decimal("0"))
        for row in rows:
            data[row["periodo"]] += row["total"] or Decimal("0")
        refund_qs = models.Devoluciones.objects.filter(
            venta__creditoshistorialcompras__isnull=True
        )
        if start is not None:
            refund_qs = refund_qs.filter(fecha__gte=start)
        if end is not None:
            refund_qs = refund_qs.filter(fecha__lt=end)
        refund_qs = _refund_condition_filter(refund_qs, condition)
        refund_rows = (
            refund_qs.annotate(periodo=period_expr)
            .values("periodo")
            .annotate(total=Coalesce(Sum("ingreso_afectado"), Decimal("0")))
        )
        for row in refund_rows:
            data[row["periodo"]] -= row["total"] or Decimal("0")
        return data

    def _sum_cash_range(qs, start=None, end=None, condition=None):
        filtered = qs
        if start is not None:
            filtered = filtered.filter(venta__fecha__gte=start)
        if end is not None:
            filtered = filtered.filter(venta__fecha__lt=end)
        total = filtered.aggregate(total=Coalesce(Sum("subtotal"), Decimal("0")))[
            "total"
        ] or Decimal("0")
        refund_qs = models.Devoluciones.objects.filter(
            venta__creditoshistorialcompras__isnull=True
        )
        if start is not None:
            refund_qs = refund_qs.filter(fecha__gte=start)
        if end is not None:
            refund_qs = refund_qs.filter(fecha__lt=end)
        refund_qs = _refund_condition_filter(refund_qs, condition)
        refund_total = refund_qs.aggregate(
            total=Coalesce(Sum("ingreso_afectado"), Decimal("0"))
        )["total"] or Decimal("0")
        return total - refund_total

    def _aggregate_credit_period(period_expr, start=None, end=None):
        qs = models.PagosCredito.objects.all()
        if start is not None:
            qs = qs.filter(fecha__gte=start)
        if end is not None:
            qs = qs.filter(fecha__lt=end)
        rows = (
            qs.annotate(periodo=period_expr)
            .values("periodo", "credito_id")
            .annotate(total=Coalesce(Sum("monto"), Decimal("0")))
        )
        new_map = defaultdict(lambda: Decimal("0"))
        used_map = defaultdict(lambda: Decimal("0"))
        for row in rows:
            ratio = credit_ratios.get(row["credito_id"], default_ratio)
            total = row["total"] or Decimal("0")
            new_map[row["periodo"]] += total * ratio["new"]
            used_map[row["periodo"]] += total * ratio["used"]
        refund_qs = models.Devoluciones.objects.filter(
            venta__creditoshistorialcompras__isnull=False
        )
        if start is not None:
            refund_qs = refund_qs.filter(fecha__gte=start)
        if end is not None:
            refund_qs = refund_qs.filter(fecha__lt=end)
        refund_rows = (
            refund_qs.annotate(periodo=period_expr)
            .values("periodo", "producto_condicion_snapshot", "producto__condicion")
            .annotate(total=Coalesce(Sum("ingreso_afectado"), Decimal("0")))
        )
        for row in refund_rows:
            cond = row["producto_condicion_snapshot"] or row["producto__condicion"] or ""
            cond = cond.lower()
            if cond == "used":
                used_map[row["periodo"]] -= row["total"] or Decimal("0")
            else:
                new_map[row["periodo"]] -= row["total"] or Decimal("0")
        return new_map, used_map

    def _sum_credit_range(start=None, end=None):
        qs = models.PagosCredito.objects.all()
        if start is not None:
            qs = qs.filter(fecha__gte=start)
        if end is not None:
            qs = qs.filter(fecha__lt=end)
        rows = qs.values("credito_id").annotate(total=Coalesce(Sum("monto"), Decimal("0")))
        total_new = Decimal("0")
        total_used = Decimal("0")
        for row in rows:
            ratio = credit_ratios.get(row["credito_id"], default_ratio)
            amount = row["total"] or Decimal("0")
            total_new += amount * ratio["new"]
            total_used += amount * ratio["used"]
        refund_qs = models.Devoluciones.objects.filter(
            venta__creditoshistorialcompras__isnull=False
        )
        if start is not None:
            refund_qs = refund_qs.filter(fecha__gte=start)
        if end is not None:
            refund_qs = refund_qs.filter(fecha__lt=end)
        refund_rows = refund_qs.values(
            "producto_condicion_snapshot", "producto__condicion"
        ).annotate(total=Coalesce(Sum("ingreso_afectado"), Decimal("0")))
        refund_new = Decimal("0")
        refund_used = Decimal("0")
        for row in refund_rows:
            cond = row["producto_condicion_snapshot"] or row["producto__condicion"] or ""
            cond = cond.lower()
            if cond == "used":
                refund_used += row["total"] or Decimal("0")
            else:
                refund_new += row["total"] or Decimal("0")
        return total_new - refund_new, total_used - refund_used

    cash_new_qs = cash_details_base.filter(condition_new)
    cash_used_qs = cash_details_base.filter(condition_used)

    cash_today_new = _sum_cash_range(cash_new_qs, start_today, end_today, condition="new")
    cash_today_used = _sum_cash_range(cash_used_qs, start_today, end_today, condition="used")
    credit_today_new, credit_today_used = _sum_credit_range(
        start=start_today, end=end_today
    )

    if section == "new":
        ventas_hoy = cash_today_new + credit_today_new
    elif section == "used":
        ventas_hoy = cash_today_used + credit_today_used
    else:
        ventas_hoy = (
            cash_today_new
            + cash_today_used
            + credit_today_new
            + credit_today_used
        )
    stats["ventas_hoy"] = ventas_hoy

    # Diario: últimos 7 días
    start_day = today - timedelta(days=6)
    start_day_dt = _aware_datetime(start_day)
    diario_profit_qs = (
        detalles_profit.filter(venta__fecha__date__gte=start_day)
        .annotate(periodo=TruncDay("venta__fecha", tzinfo=tz))
        .values("periodo")
        .annotate(
            price_total=Coalesce(Sum("subtotal"), Decimal("0")),
            cost_total=Coalesce(Sum(cost_expr), Decimal("0")),
        )
        .order_by("periodo")
    )
    profit_map = {d["periodo"]: d for d in diario_profit_qs}

    diario_cash_new = _aggregate_cash_period(
        cash_new_qs,
        TruncDay("venta__fecha", tzinfo=tz),
        start=start_day_dt,
        condition="new",
    )
    diario_cash_used = _aggregate_cash_period(
        cash_used_qs,
        TruncDay("venta__fecha", tzinfo=tz),
        start=start_day_dt,
        condition="used",
    )
    diario_credit_new, diario_credit_used = _aggregate_credit_period(
        TruncDay("fecha", tzinfo=tz), start=start_day_dt
    )

    diario = []
    for offset in range(7):
        current_day = start_day + timedelta(days=offset)
        periodo_dt = _aware_datetime(current_day)
        new_total = diario_cash_new.get(periodo_dt, Decimal("0")) + diario_credit_new.get(
            periodo_dt, Decimal("0")
        )
        used_total = diario_cash_used.get(periodo_dt, Decimal("0")) + diario_credit_used.get(
            periodo_dt, Decimal("0")
        )
        if section == "new":
            ventas_total = new_total
        elif section == "used":
            ventas_total = used_total
        else:
            ventas_total = new_total + used_total
        profit_data = profit_map.get(
            periodo_dt, {"price_total": Decimal("0"), "cost_total": Decimal("0")}
        )
        profit = profit_data["price_total"] - profit_data["cost_total"]
        diario.append(
            {
                "periodo": periodo_dt.strftime("%Y-%m-%d"),
                "ventas": ventas_total,
                "utilidad": profit,
                "price_total": profit_data["price_total"],
                "cost_total": profit_data["cost_total"],
                "profit_total": profit,
            }
        )

    # Quincenal: este mes
    year = today.year
    month = today.month
    last_day = calendar.monthrange(year, month)[1]
    quincenal = []
    first_half_profit = detalles_profit.filter(
        venta__fecha__year=year, venta__fecha__month=month, venta__fecha__day__lte=15
    ).aggregate(
        price_total=Coalesce(Sum("subtotal"), Decimal("0")),
        cost_total=Coalesce(Sum(cost_expr), Decimal("0")),
    )
    second_half_profit = detalles_profit.filter(
        venta__fecha__year=year, venta__fecha__month=month, venta__fecha__day__gt=15
    ).aggregate(
        price_total=Coalesce(Sum("subtotal"), Decimal("0")),
        cost_total=Coalesce(Sum(cost_expr), Decimal("0")),
    )
    month_start_dt = _aware_datetime(date(year, month, 1))
    mid_month_dt = _aware_datetime(date(year, month, 16))
    if month == 12:
        next_month = date(year + 1, 1, 1)
    else:
        next_month = date(year, month + 1, 1)
    next_month_dt = _aware_datetime(next_month)

    first_half_cash_new = _sum_cash_range(cash_new_qs, month_start_dt, mid_month_dt)
    first_half_cash_used = _sum_cash_range(cash_used_qs, month_start_dt, mid_month_dt)
    second_half_cash_new = _sum_cash_range(cash_new_qs, mid_month_dt, next_month_dt)
    second_half_cash_used = _sum_cash_range(cash_used_qs, mid_month_dt, next_month_dt)

    first_half_credit_new, first_half_credit_used = _sum_credit_range(
        start=month_start_dt, end=mid_month_dt
    )
    second_half_credit_new, second_half_credit_used = _sum_credit_range(
        start=mid_month_dt, end=next_month_dt
    )

    first_half_new = first_half_cash_new + first_half_credit_new
    first_half_used = first_half_cash_used + first_half_credit_used
    second_half_new = second_half_cash_new + second_half_credit_new
    second_half_used = second_half_cash_used + second_half_credit_used

    for label, sales_data, profit_data in (
        (
            "1-15",
            {"new": first_half_new, "used": first_half_used},
            first_half_profit,
        ),
        (
            f"16-{last_day}",
            {"new": second_half_new, "used": second_half_used},
            second_half_profit,
        ),
    ):
        if section == "new":
            ventas_total = sales_data["new"]
        elif section == "used":
            ventas_total = sales_data["used"]
        else:
            ventas_total = sales_data["new"] + sales_data["used"]
        profit = profit_data["price_total"] - profit_data["cost_total"]
        quincenal.append(
            {
                "periodo": label,
                "ventas": ventas_total,
                "utilidad": profit,
                "price_total": profit_data["price_total"],
                "cost_total": profit_data["cost_total"],
                "profit_total": profit,
            }
        )

    # Mensual: este año
    mensual_profit_qs = (
        detalles_profit.filter(venta__fecha__year=today.year)
        .annotate(periodo=TruncMonth("venta__fecha", tzinfo=tz))
        .values("periodo")
        .annotate(
            price_total=Coalesce(Sum("subtotal"), Decimal("0")),
            cost_total=Coalesce(Sum(cost_expr), Decimal("0")),
        )
        .order_by("periodo")
    )
    mensual = []
    profit_map = {d["periodo"]: d for d in mensual_profit_qs}
    year_start_dt = _aware_datetime(date(today.year, 1, 1))
    year_end_dt = _aware_datetime(date(today.year + 1, 1, 1))
    mensual_cash_new = _aggregate_cash_period(
        cash_new_qs,
        TruncMonth("venta__fecha", tzinfo=tz),
        start=year_start_dt,
        end=year_end_dt,
        condition="new",
    )
    mensual_cash_used = _aggregate_cash_period(
        cash_used_qs,
        TruncMonth("venta__fecha", tzinfo=tz),
        start=year_start_dt,
        end=year_end_dt,
        condition="used",
    )
    mensual_credit_new, mensual_credit_used = _aggregate_credit_period(
        TruncMonth("fecha", tzinfo=tz), start=year_start_dt, end=year_end_dt
    )
    all_month_keys = sorted(
        set(mensual_cash_new)
        | set(mensual_cash_used)
        | set(mensual_credit_new)
        | set(mensual_credit_used)
        | set(profit_map)
    )
    for periodo in all_month_keys:
        new_total = mensual_cash_new.get(periodo, Decimal("0")) + mensual_credit_new.get(
            periodo, Decimal("0")
        )
        used_total = mensual_cash_used.get(periodo, Decimal("0")) + mensual_credit_used.get(
            periodo, Decimal("0")
        )
        if section == "new":
            sales_total = new_total
        elif section == "used":
            sales_total = used_total
        else:
            sales_total = new_total + used_total
        profit_data = profit_map.get(
            periodo, {"price_total": Decimal("0"), "cost_total": Decimal("0")}
        )
        profit = profit_data["price_total"] - profit_data["cost_total"]
        mensual.append(
            {
                "periodo": periodo.strftime("%b"),
                "ventas": sales_total,
                "utilidad": profit,
                "price_total": profit_data["price_total"],
                "cost_total": profit_data["cost_total"],
                "profit_total": profit,
            }
        )

    # Todos: por año
    todos_profit_qs = (
        detalles_profit.annotate(periodo=TruncYear("venta__fecha", tzinfo=tz))
        .values("periodo")
        .annotate(
            price_total=Coalesce(Sum("subtotal"), Decimal("0")),
            cost_total=Coalesce(Sum(cost_expr), Decimal("0")),
        )
        .order_by("periodo")
    )
    todos = []
    profit_map = {d["periodo"]: d for d in todos_profit_qs}
    todos_cash_new = _aggregate_cash_period(
        cash_new_qs, TruncYear("venta__fecha", tzinfo=tz), condition="new"
    )
    todos_cash_used = _aggregate_cash_period(
        cash_used_qs, TruncYear("venta__fecha", tzinfo=tz), condition="used"
    )
    todos_credit_new, todos_credit_used = _aggregate_credit_period(
        TruncYear("fecha", tzinfo=tz)
    )
    all_year_keys = sorted(
        set(todos_cash_new)
        | set(todos_cash_used)
        | set(todos_credit_new)
        | set(todos_credit_used)
        | set(profit_map)
    )
    for periodo in all_year_keys:
        new_total = todos_cash_new.get(periodo, Decimal("0")) + todos_credit_new.get(
            periodo, Decimal("0")
        )
        used_total = todos_cash_used.get(periodo, Decimal("0")) + todos_credit_used.get(
            periodo, Decimal("0")
        )
        if section == "new":
            sales_total = new_total
        elif section == "used":
            sales_total = used_total
        else:
            sales_total = new_total + used_total
        profit_data = profit_map.get(
            periodo, {"price_total": Decimal("0"), "cost_total": Decimal("0")}
        )
        profit = profit_data["price_total"] - profit_data["cost_total"]
        todos.append(
            {
                "periodo": periodo.year,
                "ventas": sales_total,
                "utilidad": profit,
                "price_total": profit_data["price_total"],
                "cost_total": profit_data["cost_total"],
                "profit_total": profit,
            }
        )

    sales_chart = {
        "diario": diario,
        "quincenal": quincenal,
        "mensual": mensual,
        "todos": todos,
    }

    # Recent sales
    ventas_qs = models.Ventas.objects.select_related("cliente")
    if prod_cond:
        ventas_qs = ventas_qs.filter(detalles__producto__condicion=prod_cond).distinct()
    recent_sales = []
    for v in ventas_qs.order_by("-fecha")[:5]:
        customer = (
            v.cliente.razon_social
            or v.cliente.nombre
            if v.cliente
            else "Sin cliente"
        )
        recent_sales.append(
            {
                "id": v.id,
                "customer": customer,
                "amount": v.total,
                "status": v.estado,
            }
        )

    # Low stock items
    if section == "all":
        low_stock_qs = models.Productos.objects.filter(
            stock__lte=F("stock_minimo"), condicion="new"
        ).order_by("stock")[:5]
    else:
        low_stock_qs = models.Productos.objects.filter(
            stock__lte=F("stock_minimo"), **product_filter
        ).order_by("stock")[:5]
    low_stock_items = [
        {"name": p.nombre, "stock": p.stock, "min": p.stock_minimo}
        for p in low_stock_qs
    ]

    # Category data
    if prod_cond:
        category_qs = (
            models.Categorias.objects.annotate(
                value=Count("productos", filter=Q(productos__condicion=prod_cond))
            )
            .filter(value__gt=0)
            .order_by("-value")
        )
    else:
        category_qs = (
            models.Categorias.objects.annotate(value=Count("productos"))
            .filter(value__gt=0)
            .order_by("-value")
        )
    category_data = [
        {"name": c.nombre, "value": c.value} for c in category_qs
    ]

    # Top products
    base_detalles_top = detalles_sales
    if section == "all":
        base_detalles_top = detalles_sales.filter(producto__condicion="new")
    top_qs = (
        base_detalles_top.values(
            "producto_id", "producto_nombre_snapshot", "producto__nombre"
        )
        .annotate(
            ventas=Coalesce(Sum("cantidad"), Decimal("0")),
            ingresos=Coalesce(Sum("subtotal"), Decimal("0")),
        )
        .order_by("-ventas")[:5]
    )
    top_products = [
        {
            "nombre": t["producto_nombre_snapshot"] or t["producto__nombre"],
            "ventas": t["ventas"],
            "ingresos": t["ingresos"],
        }
        for t in top_qs
    ]

    data = {
        "stats": stats,
        "sales_chart": sales_chart,
        "recent_sales": recent_sales,
        "low_stock_items": low_stock_items,
        "category_data": category_data,
        "top_products": top_products,
    }

    return Response(data)


class ClientesViewSet(viewsets.ModelViewSet):
    serializer_class = serializers.ClientesSerializer
    permission_classes = [AllowAny]
    queryset = models.Clientes.objects.all()

    def get_queryset(self):
        try:
            qs = (
                self.queryset.annotate(
                    ultima_compra=Coalesce("fecha_ultima_compra", Max("ventas__fecha"))
                ).order_by("-id")
            )
            q = self.request.query_params.get("q")
            if q:
                qs = qs.filter(
                    Q(nombre__icontains=q)
                    | Q(razon_social__icontains=q)
                    | Q(nombre_comercial__icontains=q)
                    | Q(dui__icontains=q)
                    | Q(nit__icontains=q)
                    | Q(nrc__icontains=q)
                    | Q(email__icontains=q)
                    | Q(telefono__icontains=q)
                    | Q(departamento__icontains=q)
                    | Q(municipio__icontains=q)
                )
            return qs
        except Exception as e:
            import traceback
            error_detail = str(e)
            traceback.print_exc()
            # Retornar queryset vacío en caso de error
            return self.queryset.none()
    
    def list(self, request, *args, **kwargs):
        try:
            return super().list(request, *args, **kwargs)
        except Exception as e:
            import traceback
            error_detail = str(e)
            traceback.print_exc()
            return Response({"detail": f"Server error: {error_detail}"}, status=500)

    def create(self, request, *args, **kw):
        s = serializers.ClientesSerializer(data=request.data)
        try:
            s.is_valid(raise_exception=True)
            with transaction.atomic():
                obj = s.save()
            return Response(serializers.ClientesSerializer(obj).data, status=201)
        except IntegrityError as e:
            msg = str(e.__cause__ or e)
            if "clientes_nit_key" in msg:
                return Response({"nit": "NIT ya existe."}, status=400)
            if "clientes_nrc_key" in msg:
                return Response({"nrc": "NRC ya existe."}, status=400)
            return Response({"detail": msg[:200]}, status=400)


class CategoriasViewSet(viewsets.ModelViewSet):
    queryset = models.Categorias.objects.all()
    serializer_class = serializers.CategoriaSerializer
    permission_classes = [AllowAny]

    def get_queryset(self):
        qs = models.Categorias.objects.annotate(
            product_count=Count('productos')
        ).order_by('nombre')
        search = self.request.query_params.get('search')
        if search:
            qs = qs.filter(nombre__icontains=search)
        return qs

    def list(self, request, *args, **kwargs):
        qs = self.get_queryset()
        try:
            limit = int(request.query_params.get('limit', 0))
        except (TypeError, ValueError):
            limit = 0
        try:
            offset = int(request.query_params.get('offset', 0))
        except (TypeError, ValueError):
            offset = 0
        if offset:
            qs = qs[offset:]
        if limit:
            qs = qs[:limit]
        s = self.get_serializer(qs, many=True)
        return Response(s.data)

    def destroy(self, request, *args, **kwargs):
        instance = self.get_object()
        if instance.productos.exists():
            return Response(
                {"detail": "La categoría tiene productos"}, status=409
            )
        return super().destroy(request, *args, **kwargs)


class ProductosViewSet(viewsets.ModelViewSet):
    queryset = models.Productos.objects.select_related("categoria").all()
    serializer_class = serializers.ProductosSerializer
    permission_classes = [IsAuthenticated]

    def _get_role(self, request):
        perfil = getattr(request.user, "usuario", None)
        return getattr(perfil, "role", None) or "vendedor"

    def _has_inventory_restrictions(self, request):
        # Sin inventario, no hay restricciones de inventario
        return False

    def _is_read_only_vendor(self, request):
        # En facturador simplificado, todos pueden editar
        return False

    def get_queryset(self):
        qs = (
            models.Productos.objects.select_related("categoria")
            .only(
                "id",
                "codigo",
                "nombre",
                "precio",
                "tipo",
                "status",
                "categoria_id",
                "categoria__nombre",
                "created_at",
                "updated_at",
            )
            .all()
        )

        tipo = self.request.query_params.get("tipo")
        status = self.request.query_params.get("status")
        categoria_param = self.request.query_params.get("categoria_id")
        search = (self.request.query_params.get("q") or self.request.query_params.get("search") or "").strip()
        ordering = (self.request.query_params.get("ordering") or "").strip()

        if tipo in ("producto", "servicio"):
            qs = qs.filter(tipo=tipo)
        if status in ("active", "archived"):
            qs = qs.filter(status=status)

        if categoria_param:
            ids = [int(pk) for pk in categoria_param.split(",") if pk.isdigit()]
            if ids:
                if len(ids) == 1:
                    qs = qs.filter(categoria_id=ids[0])
                else:
                    qs = qs.filter(categoria_id__in=ids)

        if search:
            qs = qs.filter(Q(nombre__icontains=search) | Q(codigo__icontains=search))

        allowed_ordering = {
            "nombre": "nombre",
            "-nombre": "-nombre",
            "precio": "precio",
            "-precio": "-precio",
            "codigo": "codigo",
            "-codigo": "-codigo",
            "tipo": "tipo",
            "-tipo": "-tipo",
            "updated_at": "updated_at",
            "-updated_at": "-updated_at",
        }
        default_order = "-updated_at"
        order = allowed_ordering.get(ordering, default_order)
        ordering_fields = [order]
        if order != "-updated_at":
            ordering_fields.append("-updated_at")
        ordering_fields.append("id")
        return qs.order_by(*ordering_fields)

    def list(self, request, *args, **kwargs):
        try:
            queryset = self.filter_queryset(self.get_queryset())

            full_path = request.get_full_path()
            user_segment = getattr(request.user, "pk", "anon")
            cache_key = f"productos:list:{user_segment}:{full_path}"
            cached_payload = cache.get(cache_key)
            if cached_payload is not None:
                return Response(cached_payload)

            try:
                page = int(request.query_params.get("page", 1))
            except (TypeError, ValueError):
                page = 1
            if page < 1:
                page = 1

            try:
                page_size = int(request.query_params.get("page_size", 50))
            except (TypeError, ValueError):
                page_size = 50
            page_size = max(1, min(page_size, 100))

            with transaction.atomic():
                with connection.cursor() as cursor:
                    cursor.execute("SET LOCAL statement_timeout TO %s", [10000])
                total_count = queryset.count()
                offset = (page - 1) * page_size
                results = list(queryset[offset : offset + page_size])

            serializer = self.get_serializer(results, many=True)
            total_pages = math.ceil(total_count / page_size) if total_count else 0

            base_url = request.build_absolute_uri(request.path)

            def build_page_url(target_page):
                if not target_page:
                    return None
                params = request.query_params.copy()
                params._mutable = True
                if target_page == 1:
                    params.pop("page", None)
                else:
                    params["page"] = str(target_page)
                query = params.urlencode()
                return f"{base_url}{'?' + query if query else ''}"

            next_page = page + 1 if offset + page_size < total_count else None
            previous_page = page - 1 if page > 1 else None

            payload = {
                "count": total_count,
                "page": page,
                "page_size": page_size,
                "total_pages": total_pages,
                "next": build_page_url(next_page),
                "previous": build_page_url(previous_page),
                "results": serializer.data,
            }
            cache.set(cache_key, payload, 30)
            return Response(payload)
        except Exception as e:
            import traceback
            error_detail = str(e)
            traceback.print_exc()
            return Response({"detail": f"Server error: {error_detail}"}, status=500)

    def _check_vendor_changes(self, request, instance):
        # Simplificado: solo verificar precio si es necesario
        data = request.data
        if not isinstance(data, dict):
            data = getattr(data, "dict", lambda: {})(mutable=False)
        if not data:
            return None
        # En un facturador simplificado, los vendedores pueden editar precios
        return None

    def create(self, request, *args, **kwargs):
        if self._is_read_only_vendor(request):
            return Response({"detail": "Permiso requerido"}, status=403)
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        obj = serializer.save()
        return Response(self.get_serializer(obj).data, status=201)

    def update(self, request, *args, **kwargs):
        instance = self.get_object()
        if self._is_read_only_vendor(request):
            return Response({"detail": "Permiso requerido"}, status=403)
        if self._has_inventory_restrictions(request):
            changed = self._check_vendor_changes(request, instance)
            if changed:
                return Response(
                    {changed: "Permiso requerido"},
                    status=422,
                )
        return super().update(request, *args, **kwargs)

    def partial_update(self, request, *args, **kwargs):
        instance = self.get_object()
        if self._is_read_only_vendor(request):
            return Response({"detail": "Permiso requerido"}, status=403)
        if self._has_inventory_restrictions(request):
            changed = self._check_vendor_changes(request, instance)
            if changed:
                return Response(
                    {changed: "Permiso requerido"},
                    status=422,
                )
        return super().partial_update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        if self._has_inventory_restrictions(request):
            return Response({"detail": "Permiso requerido"}, status=403)
        return super().destroy(request, *args, **kwargs)


class VentasViewSet(viewsets.ModelViewSet):
    queryset = models.Ventas.objects.select_related("cliente").all()
    serializer_class = serializers.VentasSerializer

    def get_serializer_class(self):
        if self.action == "retrieve":
            return serializers.VentaDetalleSerializer
        return super().get_serializer_class()


class VentaItemsAPIView(ListAPIView):
    serializer_class = serializers.VentaItemSerializer
    permission_classes = [AllowAny]

    def get_queryset(self):
        return models.DetalleVenta.objects.filter(venta_id=self.kwargs["pk"]).select_related("producto")


class DetalleVentaViewSet(viewsets.ModelViewSet):
    queryset = models.DetalleVenta.objects.all()
    serializer_class = serializers.DetalleVentaSerializer


class CreditosViewSet(viewsets.ModelViewSet):
    queryset = (
        models.Creditos.objects.select_related("cliente")
        .prefetch_related("pagos", "historial__venta__detalles__producto")
        .all()
    )
    serializer_class = serializers.CreditosSerializer

    def get_serializer_class(self):
        if self.action == "retrieve":
            return serializers.CreditoDetalleSerializer
        return super().get_serializer_class()

    @action(detail=True, methods=["get"])
    def pagos(self, request, pk=None):
        qs = models.PagosCredito.objects.filter(credito_id=pk).order_by("-fecha")
        return Response(serializers.PagosCreditoSerializer(qs, many=True).data)

    @action(detail=True, methods=["get"])
    def items(self, request, pk=None):
        detalles = models.DetalleVenta.objects.filter(
            venta__creditoshistorialcompras__credito_id=pk
        ).select_related("producto")
        data = [
            {
                "producto_id": d.producto_id,
                "codigo": d.producto_codigo_snapshot or getattr(d.producto, "codigo", ""),
                "nombre": d.producto_nombre_snapshot or getattr(d.producto, "nombre", ""),
                "cantidad": d.cantidad,
                "precio": d.precio_unitario,
                "sub_total": d.subtotal,
            }
            for d in detalles
        ]
        return Response(data)


class CreditosHistorialComprasViewSet(viewsets.ModelViewSet):
    queryset = models.CreditosHistorialCompras.objects.all()
    serializer_class = serializers.CreditosHistorialComprasSerializer


class PagosCreditoViewSet(viewsets.ModelViewSet):
    queryset = models.PagosCredito.objects.all()
    serializer_class = serializers.PagosCreditoSerializer


class DeudoresListAPIView(ListAPIView):
    permission_classes = [AllowAny]
    serializer_class = serializers.DeudorSerializer

    def get_queryset(self):
        qs = (
            models.Creditos.objects.values('cliente_id')
            .annotate(
                cliente_nombre=F('cliente__nombre'),
                total=Coalesce(Sum('total_deuda'), Decimal('0')),
                pagado=Coalesce(Sum('pagado'), Decimal('0')),
                saldo=Coalesce(Sum('saldo'), Decimal('0')),
            )
            .order_by('cliente_nombre')
        )
        return qs


class DeudorDetailAPIView(viewsets.ViewSet):
    permission_classes = [AllowAny]

    def list(self, request):
        return Response([])

    def retrieve(self, request, pk=None):
        qs = models.Creditos.objects.filter(cliente_id=pk).prefetch_related(
            'pagos', 'historial__venta__detalles__producto', 'cliente'
        )
        if not qs.exists():
            return Response({'detail': 'No encontrado'}, status=404)
        total = sum((c.total_deuda for c in qs), Decimal('0'))
        pagos = []
        items = []
        creditos = []
        pagado_total = Decimal('0')
        for c in qs:
            pagado_c = sum((p.monto for p in c.pagos.all()), Decimal('0'))
            pagado_total += pagado_c
            creditos.append({
                'id': c.id,
                'fecha': c.fecha_ultima_compra,
                'total': c.total_deuda,
                'pagado': pagado_c,
                'saldo': c.total_deuda - pagado_c,
                'observaciones': c.observaciones,
            })
            for p in c.pagos.all():
                pagos.append({'id': p.id, 'fecha': p.fecha, 'monto': p.monto})
            for h in c.historial.all():
                venta = h.venta
                for d in venta.detalles.all():
                    items.append({
                        'producto_id': d.producto_id,
                        'codigo': d.producto_codigo_snapshot or getattr(d.producto, 'codigo', ''),
                        'nombre': d.producto_nombre_snapshot or getattr(d.producto, 'nombre', ''),
                        'cantidad': d.cantidad,
                        'precio': d.precio_unitario,
                        'sub_total': d.subtotal,
                    })
        data = {
            'cliente': {'id': pk, 'nombre': qs[0].cliente.nombre},
            'total': total,
            'pagado': pagado_total,
            'saldo': total - pagado_total,
            'creditos': creditos,
            'pagos': pagos,
            'items': items,
        }
        s = serializers.DeudorDetalleSerializer(data)
        return Response(s.data)


@api_view(["GET"])
def devoluciones_search_lines(request):
    term = request.query_params.get("term", "").strip()
    if not term:
        return Response({"results": [], "count": 0})

    limit = int(request.query_params.get("limit", 50))
    offset = int(request.query_params.get("offset", 0))
    if limit > 100:
        limit = 100

    qs = (
        models.DetalleVenta.objects.select_related("venta", "venta__cliente")
        .annotate(
            qty_devuelta=Coalesce(Sum("devoluciones__cantidad"), Decimal("0"))
        )
        .filter(
            Q(producto_nombre_snapshot__icontains=term)
            | Q(producto_codigo_snapshot__icontains=term)
        )
        .order_by("-venta__fecha")
    )

    total = qs.count()
    qs = qs[offset : offset + limit]
    data = []
    for dv in qs:
        v = dv.venta
        data.append(
            {
                "detalle_venta_id": dv.id,
                "venta_id": v.id if v else None,
                "fecha": v.fecha if v else None,
                "cliente": v.cliente.nombre if v and v.cliente else None,
                "qty_vendida": dv.cantidad,
                "qty_devuelta": dv.qty_devuelta,
                "precio_unitario": dv.precio_unitario,
                "codigo_s": dv.producto_codigo_snapshot,
                "nombre_s": dv.producto_nombre_snapshot,
                "condicion_s": dv.producto_condicion_snapshot,
                "categoria_s": dv.producto_categoria_nombre_snapshot,
            }
        )

    return Response({"results": data, "count": total})


class DevolucionesViewSet(viewsets.ModelViewSet):
    queryset = models.Devoluciones.objects.all()
    serializer_class = serializers.DevolucionesSerializer

    def create(self, request, *args, **kwargs):
        venta_id = request.data.get("venta_id")
        items = request.data.get("items") or []
        if not venta_id:
            return Response({"detail": "venta_id requerido"}, status=400)
        if not isinstance(items, list) or not items:
            return Response({"detail": "Se requieren items"}, status=400)

        now = timezone.now()

        try:
            with transaction.atomic():
                try:
                    venta = models.Ventas.objects.select_for_update().get(
                        id=venta_id
                    )
                except models.Ventas.DoesNotExist:
                    return Response({"detail": "Venta no encontrada"}, status=404)

                parsed_items = []
                detalle_ids = []
                for raw in items:
                    detalle_id = raw.get("detalle_id") or raw.get("detalle")
                    qty_raw = raw.get("qty") or raw.get("cantidad")
                    motivo = raw.get("motivo")
                    if not detalle_id:
                        return Response({"detail": "detalle_id requerido"}, status=400)
                    try:
                        qty = Decimal(str(qty_raw))
                    except Exception:
                        return Response({"detail": "Cantidad inválida"}, status=400)
                    qty = qty.quantize(Decimal("0.001"), rounding=ROUND_HALF_UP)
                    if qty <= 0:
                        return Response({"detail": "Cantidad debe ser positiva"}, status=400)
                    parsed_items.append({
                        "detalle_id": detalle_id,
                        "qty": qty,
                        "motivo": motivo,
                    })
                    detalle_ids.append(detalle_id)

                detalles_map = {
                    det.id: det
                    for det in models.DetalleVenta.objects.select_for_update()
                    .filter(venta_id=venta_id, id__in=detalle_ids)
                }

                producto_ids = {
                    det.producto_id
                    for det in detalles_map.values()
                    if det.producto_id
                }
                productos_map = {}
                if producto_ids:
                    productos_map = {
                        prod.id: prod
                        for prod in models.Productos.objects.select_for_update()
                        .filter(id__in=producto_ids)
                    }

                credit_hist = (
                    models.CreditosHistorialCompras.objects.select_related("credito")
                    .filter(venta_id=venta_id)
                    .order_by("id")
                    .first()
                )
                credito = None
                if credit_hist:
                    credito = (
                        models.Creditos.objects.select_for_update()
                        .get(id=credit_hist.credito_id)
                    )

                processed = []
                total_refund = Decimal("0")

                for parsed in parsed_items:
                    detalle_id = parsed["detalle_id"]
                    detalle = detalles_map.get(detalle_id)
                    if not detalle:
                        return Response(
                            {
                                "detail": f"Detalle {detalle_id} no pertenece a la venta"
                            },
                            status=404,
                        )
                    qty = parsed["qty"]
                    motivo = parsed["motivo"]
                    producto = (
                        productos_map.get(detalle.producto_id)
                        if detalle.producto_id
                        else None
                    )

                    disponible = detalle.cantidad - (detalle.devuelto or Decimal("0"))
                    if qty > disponible:
                        return Response(
                            {
                                "detail": "Cantidad supera disponible",
                                "detalle_id": detalle_id,
                                "disponible": float(disponible),
                            },
                            status=400,
                        )

                    line_total = (qty * detalle.precio_unitario).quantize(
                        Decimal("0.01"), rounding=ROUND_HALF_UP
                    )
                    processed.append(
                        {
                            "detalle": detalle,
                            "producto": producto,
                            "qty": qty,
                            "motivo": motivo,
                            "total": line_total,
                        }
                    )
                    total_refund += line_total

                prev_applied = models.Devoluciones.objects.filter(
                    venta_id=venta_id
                ).aggregate(total=Coalesce(Sum("ingreso_afectado"), Decimal("0")))[
                    "total"
                ] or Decimal("0")

                income_to_allocate = total_refund
                pagos_total = Decimal("0")
                if credito:
                    pagos_total = (
                        models.PagosCredito.objects.filter(credito_id=credito.id)
                        .aggregate(total=Coalesce(Sum("monto"), Decimal("0")))["total"]
                        or Decimal("0")
                    )
                    disponible_ingreso = max(pagos_total - prev_applied, Decimal("0"))
                    income_to_allocate = min(disponible_ingreso, total_refund)

                allocated = []
                remaining_income = income_to_allocate
                item_count = len(processed)
                for idx, item in enumerate(processed):
                    detalle = item["detalle"]
                    line_income = Decimal("0")
                    if income_to_allocate > 0:
                        if credito:
                            if item_count == 1 or idx == item_count - 1:
                                line_income = remaining_income
                            elif total_refund > 0:
                                line_income = (
                                    income_to_allocate
                                    * item["total"]
                                    / total_refund
                                ).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
                            remaining_income -= line_income
                            if remaining_income < Decimal("0"):
                                remaining_income = Decimal("0")
                        else:
                            line_income = item["total"]

                    allocated.append(
                        {
                            "detalle": detalle,
                            "producto": item["producto"],
                            "qty": item["qty"],
                            "motivo": item["motivo"],
                            "total": item["total"],
                            "ingreso": max(line_income, Decimal("0")),
                        }
                    )

                created = []
                for info in allocated:
                    detalle = info["detalle"]
                    qty = info["qty"]
                    motivo = info["motivo"]
                    line_total = info["total"]
                    ingreso = info["ingreso"]
                    producto = info.get("producto")

                    producto_id = detalle.producto_id if detalle.producto_id else None

                    devolucion = models.Devoluciones.objects.create(
                        fecha=now,
                        producto_id=producto_id,
                        venta_id=venta_id,
                        detalle_venta_id=detalle.id,
                        cantidad=qty,
                        precio_unitario=detalle.precio_unitario,
                        total=line_total,
                        motivo=motivo,
                        ingreso_afectado=ingreso,
                        producto_codigo_snapshot=detalle.producto_codigo_snapshot
                        or (producto.codigo if producto else None),
                        producto_nombre_snapshot=detalle.producto_nombre_snapshot
                        or (producto.nombre if producto else None),
                        producto_costo_snapshot=detalle.producto_costo_snapshot
                        or (producto.costo if producto else None),
                        producto_condicion_snapshot=detalle.producto_condicion_snapshot
                        or (producto.condicion if producto else None),
                        producto_categoria_id_snapshot=detalle.producto_categoria_id_snapshot
                        or (
                            producto.categoria_id if producto else None
                        ),
                        producto_categoria_nombre_snapshot=detalle.producto_categoria_nombre_snapshot
                        or (
                            producto.categoria.nombre
                            if producto and producto.categoria
                            else None
                        ),
                        created_at=now,
                        updated_at=now,
                    )
                    created.append(devolucion)

                    models.DetalleVenta.objects.filter(id=detalle.id).update(
                        devuelto=F("devuelto") + qty,
                        updated_at=now,
                    )

                    if producto_id:
                        models.Productos.objects.filter(id=producto_id).update(
                            stock=F("stock") + qty,
                            updated_at=now,
                        )

                if credito:
                    nuevo_total = max(credito.total_deuda - total_refund, Decimal("0"))
                    nuevo_pagado = max(credito.pagado - income_to_allocate, Decimal("0"))
                    nuevo_saldo = max(nuevo_total - nuevo_pagado, Decimal("0"))
                    credito.total_deuda = nuevo_total
                    credito.pagado = nuevo_pagado
                    credito.saldo = nuevo_saldo
                    credito.estado = "pagado" if nuevo_saldo == 0 else "pendiente"
                    credito.updated_at = now
                    credito.save(
                        update_fields=[
                            "total_deuda",
                            "pagado",
                            "saldo",
                            "estado",
                            "updated_at",
                        ]
                    )

                    models.CreditosHistorialCompras.objects.create(
                        credito=credito,
                        venta_id=venta_id,
                        fecha=now,
                        monto=-total_refund,
                        pagado=-income_to_allocate,
                        saldo=credito.saldo,
                        estado=credito.estado,
                        created_at=now,
                        updated_at=now,
                    )

                return Response(
                    {
                        "ok": True,
                        "venta_id": venta_id,
                        "total_refund": float(total_refund),
                        "ingreso_afectado": float(income_to_allocate),
                        "items": [
                            {
                                "devolucion_id": dev.id,
                                "detalle_id": dev.detalle_venta_id,
                                "cantidad": float(dev.cantidad),
                                "total": float(dev.total),
                            }
                            for dev in created
                        ],
                    },
                    status=201,
                )
        except DataError as exc:
            transaction.set_rollback(True)
            return Response({"detail": str(exc)}, status=400)



class ValidateOverrideCodeView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        ip = request.META.get("REMOTE_ADDR", "")
        key = f"ovr:{ip}"
        ttl = settings.RATE_LIMIT_OVERRIDE_TTL
        max_attempts = settings.RATE_LIMIT_OVERRIDE_MAX_ATTEMPTS
        data = cache.get(key, {"attempts": 0, "until": 0})
        attempts = data.get("attempts", 0)
        until = data.get("until", 0)
        now = int(time.time())
        if until and until > now:
            retry_after = until - now
            return Response({"locked": True, "attempts_left": 0, "retry_after": retry_after})
        return Response({"locked": False, "attempts_left": max_attempts - attempts})

    def post(self, request):
        code = request.data.get("code", "")
        ip = request.META.get("REMOTE_ADDR", "")
        key = f"ovr:{ip}"
        ttl = settings.RATE_LIMIT_OVERRIDE_TTL
        max_attempts = settings.RATE_LIMIT_OVERRIDE_MAX_ATTEMPTS
        data = cache.get(key, {"attempts": 0, "until": 0})
        attempts = data.get("attempts", 0)
        until = data.get("until", 0)
        now = int(time.time())
        if until and until > now:
            retry_after = until - now
            return Response(
                {"ok": False, "locked": True, "attempts_left": 0, "retry_after": retry_after},
                status=429,
            )
        if constant_time_compare(code, settings.PRICE_OVERRIDE_CODE):
            cache.delete(key)
            return Response({"ok": True, "attempts_left": max_attempts})
        attempts += 1
        if attempts >= max_attempts:
            until = now + ttl
            cache.set(key, {"attempts": attempts, "until": until}, ttl)
            return Response(
                {"ok": False, "locked": True, "attempts_left": 0, "retry_after": ttl},
                status=429,
            )
        cache.set(key, {"attempts": attempts, "until": 0}, ttl)
        return Response(
            {"ok": False, "attempts_left": max_attempts - attempts, "retry_after": ttl},
            status=401,
        )


@api_view(["POST"])
@parser_classes([JSONParser])
def pos_checkout(request):
    d = request.data
    now = timezone.now()
    is_credit = d["saleType"] == "CREDIT"
    tot = Decimal(str(d["totals"]["total"]))
    paid = Decimal(str(d.get("paidAmount", 0)))
    cust_id = d.get("customerId") or None

    if cust_id and not models.Clientes.objects.filter(id=cust_id).exists():
        return Response({"detail": "Cliente no existe"}, status=400)

    method = PAYMENT_METHOD_MAP[d["paymentMethod"]]

    if is_credit:
        if not cust_id:
            return Response({"detail": "Cliente requerido para crédito"}, status=400)
        if not (Decimal("0") <= paid <= tot):
            return Response({"detail": "Abono inválido"}, status=400)
    else:
        if method == "efectivo" and paid < tot:
            return Response({"detail": "Efectivo debe cubrir total"}, status=400)
        if method != "efectivo" and paid != tot:
            return Response({"detail": "Pago debe cubrir total"}, status=400)

    items = d.get("items", [])
    if not items or tot <= 0:
        return Response({"detail": "La venta debe tener items y total mayor a 0"}, status=400)

    try:
        with transaction.atomic():
            venta = models.Ventas.objects.create(
                fecha=now,
                cliente_id=cust_id,
                total=tot,
                estado="completada",
                metodo_pago=method,
                created_at=now,
                updated_at=now,
            )

            for it in items:
                qty = Decimal(str(it["qty"]))
                price = Decimal(str(it["unit_price"]))
                override = bool(it.get("override"))
                models.DetalleVenta.objects.create(
                    venta=venta,
                    producto_id=it["productId"],
                    cantidad=qty,
                    precio_unitario=price,
                    subtotal=qty * price,
                    fecha_venta=now,
                    override=override,
                    created_at=now,
                    updated_at=now,
                )

            if is_credit:
                saldo = tot - paid
                credito = models.Creditos.objects.create(
                    cliente_id=cust_id,
                    total_deuda=tot,
                    pagado=paid,
                    saldo=saldo,
                    fecha_ultima_compra=now,
                    estado="pendiente",
                    observaciones=d.get("observaciones"),
                    created_at=now,
                    updated_at=now,
                )
                models.CreditosHistorialCompras.objects.create(
                    credito=credito,
                    venta=venta,
                    fecha=now,
                    monto=tot,
                    pagado=paid,
                    saldo=saldo,
                    estado="pendiente",
                    created_at=now,
                    updated_at=now,
                )
                if paid > 0:
                    models.PagosCredito.objects.create(
                        credito=credito,
                        fecha=now,
                        monto=paid,
                        concepto="Abono inicial",
                        metodo_pago=method,
                        created_at=now,
                        updated_at=now,
                    )

        return Response({"id": venta.id}, status=201)
    except (DataError, IntegrityError, KeyError) as exc:
        return Response({"detail": str(exc)}, status=400)


class UsuariosViewSet(viewsets.ModelViewSet):
    serializer_class = serializers.UsuarioSerializer
    permission_classes = [IsAuthenticated]
    queryset = models.Usuario.objects.select_related('user').all()

    def get_queryset(self):
        qs = self.queryset.order_by('-id')
        q = self.request.query_params.get("q")
        if q:
            qs = qs.filter(
                Q(user__username__icontains=q)
                | Q(user__email__icontains=q)
                | Q(role__icontains=q)
            )
        return qs

    def list(self, request, *args, **kwargs):
        try:
            return super().list(request, *args, **kwargs)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"detail": f"Server error: {str(e)}"}, status=500)

    def create(self, request, *args, **kwargs):
        import logging
        logger = logging.getLogger(__name__)
        logger.info(f"Creating usuario with data: {request.data}")
        
        serializer = self.get_serializer(data=request.data)
        try:
            serializer.is_valid(raise_exception=True)
            logger.info("Serializer is valid")
            with transaction.atomic():
                usuario = serializer.save()
            logger.info(f"Usuario created: {usuario.id}")
            # Asegurar que el usuario relacionado esté cargado
            usuario = models.Usuario.objects.select_related('user').get(id=usuario.id)
            response_data = serializer.to_representation(usuario)
            logger.info(f"Response data: {response_data}")
            return Response(response_data, status=201)
        except serializers.ValidationError as e:
            import traceback
            traceback.print_exc()
            logger.error(f"Validation error: {e.detail}")
            return Response(e.detail, status=400)
        except Exception as e:
            import traceback
            error_msg = str(e)
            logger.error(f"Error creating usuario: {error_msg}", exc_info=True)
            traceback.print_exc()
            return Response({"detail": error_msg}, status=500)

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        try:
            serializer.is_valid(raise_exception=True)
            with transaction.atomic():
                usuario = serializer.save()
            return Response(serializer.data)
        except serializers.ValidationError as e:
            return Response(e.detail, status=400)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"detail": str(e)}, status=500)
